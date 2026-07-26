"""
Document loaders for all supported ingestion paths.

Each loader returns a RawDocument dataclass. URL loading uses Trafilatura
for boilerplate removal — raw HTML is never passed downstream.
"""

import logging
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path

import fitz  # PyMuPDF
import trafilatura
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass
class RawDocument:
    """Raw document extracted by a loader, before normalization."""

    content: str  # extracted text (clean, no HTML for URLs)
    source_type: str  # "pdf" | "docx" | "url" | "message"
    source_ref: str  # file path, URL, or "pasted"
    metadata: dict = field(default_factory=dict)


def load_pdf(path: str | Path) -> RawDocument:
    """Load text content from a PDF file using PyMuPDF.

    Extracts text page-by-page and preserves page metadata.

    Args:
        path: Path to the PDF file.

    Returns:
        RawDocument with extracted text and metadata.

    Raises:
        FileNotFoundError: If the PDF file does not exist.
        ValueError: If the file cannot be parsed as PDF.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    logger.info("Loading PDF: %s", path)
    try:
        doc = fitz.open(str(path))
    except Exception as e:
        raise ValueError(f"Failed to open PDF '{path}': {e}") from e

    pages_text = []
    page_count = len(doc)  # capture before closing
    for page_num in range(page_count):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            pages_text.append(text)

    content = "\n\n".join(pages_text)

    if not content.strip():
        logger.warning("PDF '%s' produced no extractable text", path)

    raw_doc = RawDocument(
        content=content,
        source_type="pdf",
        source_ref=str(path),
        metadata={
            "filename": path.name,
            "page_count": page_count,
            "file_size_bytes": path.stat().st_size,
        },
    )
    doc.close()  # close AFTER building RawDocument
    return raw_doc


def load_docx(path: str | Path) -> RawDocument:
    """Load text content from a DOCX file using python-docx.

    Iterates paragraphs and preserves heading structure.

    Args:
        path: Path to the DOCX file.

    Returns:
        RawDocument with extracted text and metadata.

    Raises:
        FileNotFoundError: If the DOCX file does not exist.
        ValueError: If the file cannot be parsed as DOCX.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    logger.info("Loading DOCX: %s", path)
    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        raise ValueError(f"Failed to open DOCX '{path}': {e}") from e

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Prefix headings with markdown-style markers for structure
            if para.style and para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.split()[-1])
                    text = f"{'#' * level} {text}"
                except (ValueError, IndexError):
                    text = f"# {text}"
            paragraphs.append(text)

    content = "\n\n".join(paragraphs)

    if not content.strip():
        logger.warning("DOCX '%s' produced no extractable text", path)

    return RawDocument(
        content=content,
        source_type="docx",
        source_ref=str(path),
        metadata={
            "filename": path.name,
            "paragraph_count": len(paragraphs),
            "file_size_bytes": path.stat().st_size,
        },
    )


def load_url(url: str) -> RawDocument:
    """Load and extract main content from a URL using Trafilatura.

    Pipeline:
        URL → Trafilatura.fetch_url() → trafilatura.extract()
            → returns clean article text (title, headings, paragraphs)
            → removes navigation, ads, footers, scripts, sidebars
            → output is plain text / light markdown, NOT HTML

    This is critical for embedding quality — boilerplate removal means
    chunks contain only meaningful article content.

    Args:
        url: The URL to fetch and extract content from.

    Returns:
        RawDocument with clean extracted text.

    Raises:
        ValueError: If the URL cannot be fetched or no content extracted.
    """
    logger.info("Loading URL: %s", url)

    # Fetch the raw HTML
    downloaded = trafilatura.fetch_url(url)
    if downloaded is None:
        raise ValueError(f"Failed to fetch URL: {url}")

    # Extract main content with Trafilatura (removes boilerplate)
    content = trafilatura.extract(
        downloaded,
        include_comments=False,
        include_tables=True,
        output_format="txt",
        favor_precision=True,
    )

    if not content or not content.strip():
        raise ValueError(f"No extractable content from URL: {url}")

    # Try to extract the title separately
    metadata_result = trafilatura.extract(
        downloaded,
        output_format="json",
        include_comments=False,
    )
    title = ""
    if metadata_result:
        import json
        try:
            meta = json.loads(metadata_result)
            title = meta.get("title", "")
        except (json.JSONDecodeError, TypeError):
            pass

    return RawDocument(
        content=content,
        source_type="url",
        source_ref=url,
        metadata={
            "title": title,
            "fetch_date": datetime.now(timezone.utc).isoformat(),
            "content_length": len(content),
        },
    )


def load_message(text: str, source_label: str = "pasted") -> RawDocument:
    """Load a pasted message or chat thread as a document.

    Pass-through loader that wraps raw text with metadata.

    Args:
        text: The pasted text content.
        source_label: Label for the source (default: "pasted").

    Returns:
        RawDocument with the text content.

    Raises:
        ValueError: If the text is empty.
    """
    if not text or not text.strip():
        raise ValueError("Cannot load empty message")

    logger.info("Loading pasted message (%d chars)", len(text))
    return RawDocument(
        content=text.strip(),
        source_type="message",
        source_ref=source_label,
        metadata={
            "pasted_at": datetime.now(timezone.utc).isoformat(),
            "char_count": len(text.strip()),
            "line_count": len(text.strip().splitlines()),
        },
    )
